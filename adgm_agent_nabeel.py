"""
Enhanced ADGM Corporate Agent - Real ADGM Data Sources Integration
- Fetches real ADGM templates, checklists, and guidance documents
- Uses official ADGM data sources for compliance checking
- Maintains original modularity and efficiency
- Accepts only .docx files with inline markup
- Enhanced with real-world ADGM legal data via RAG
- Includes AI-check for ambiguous language and robust API key validation
"""
'''Made by Nabeel and ai assistants claude and gemini'''

import os
import re
import json
import zipfile
import logging
import tempfile
import uuid
import requests
import time
from datetime import datetime
from typing import List, Dict, Optional, Any
from dataclasses import dataclass, asdict
from urllib.parse import urlparse

import gradio as gr
import numpy as np
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX

# Graceful imports with fallbacks
try:
    import spacy
    from sentence_transformers import SentenceTransformer
    from openai import OpenAI
    import faiss
    NLP_MODEL = spacy.load("en_core_web_sm")
except ImportError as e:
    NLP_MODEL = None
    logging.warning(f"Optional dependency missing: {e}. Some NLP features will be disabled.")

# Graceful import for PDF parsing
try:
    import fitz  # PyMuPDF
    PDF_PARSER_ENABLED = True
except ImportError:
    PDF_PARSER_ENABLED = False
    logging.warning("PyMuPDF not found. PDF parsing will be disabled.")


logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("adgm_agent")

# -------------------------
# CONFIGURATION & REAL ADGM DATA
# -------------------------

@dataclass
class Config:
    embedding_model: str = "all-MiniLM-L6-v2"
    llm_model: str = "gpt-4o-mini"
    max_tokens: int = 1000
    temperature: float = 0.1
    chunk_size: int = 1000
    chunk_overlap: int = 200
    allowed_types: List[str] = None
    
    def __post_init__(self):
        if self.allowed_types is None:
            self.allowed_types = [".docx"]

    def adjust_for_upload_size(self, total_bytes: int):
        if total_bytes > 5 * 1024 * 1024:  # > 5 MB
            logger.info(f"Large upload detected ({total_bytes / 1024**2:.2f} MB). Using smaller chunks.")
            self.chunk_size = 750
            self.chunk_overlap = 150
        else:
            self.chunk_size = 1000
            self.chunk_overlap = 200
        logger.info(f"Using Chunk Size: {self.chunk_size}, Overlap: {self.chunk_overlap}")

# Static data sources and checklists (Expanded for clarity)
ADGM_DATA_SOURCES = {
    "company_formation": {
        "name": "Company Formation & Governance",
        "documents": {
            "incorporation_resolution_ltd": "https://assets.adgm.com/download/assets/adgm-ra-resolution-multiple-incorporate-shareholders-LTD-incorporation-v2.docx/186a12846c3911efa4e6c6223862cd87",
            "checklist_private_company": "https://www.adgm.com/documents/registration-authority/registration-and-incorporation/checklist/private-company-limited-by-guarantee-non-financial-services-20231228.pdf",
        },
    },
    "employment": {"name": "Employment & HR", "documents": {"standard_contract_2024": "https://assets.adgm.com/download/assets/ADGM+Standard+Employment+Contract+Template+-+ER+2024+(Feb+2025).docx/ee14b252edbe11efa63b12b3a30e5e3a"}},
    "data_protection": {"name": "Data Protection", "documents": {"appropriate_policy": "https://www.adgm.com/documents/office-of-data-protection/templates/adgm-dpr-2021-appropriate-policy-document.pdf"}},
    "compliance": {"name": "Compliance & Filings", "documents": {"shareholder_resolution_template": "https://assets.adgm.com/download/assets/Templates_SHReso_AmendmentArticles-v1-20220107.docx/97120d7c5af911efae4b1e183375c0b2"}},
}

ADGM_CHECKLISTS = {
    "company_incorporation": {"name": "Company Incorporation", "required_documents": ["Articles of Association", "Memorandum of Association", "Incorporation Application Form", "UBO Declaration Form", "Register of Members and Directors"]},
    "employment": {"name": "Employment & HR", "required_documents": ["Employment Contract", "Employee Handbook", "Work Permit Applications"]},
    "data_protection": {"name": "Data Protection Compliance", "required_documents": ["Appropriate Policy Document", "Privacy Notice", "Data Processing Agreement"]},
}

# -------------------------
# UTILITY FUNCTIONS
# -------------------------

def chunk_text(text: str, size: int, overlap: int) -> List[Dict]:
    if not text: return []
    chunks, start = [], 0
    while start < len(text):
        end = min(start + size, len(text))
        chunk_text = text[start:end].strip()
        if chunk_text: chunks.append({"text": chunk_text, "start": start, "end": end})
        if end == len(text): break
        start = end - overlap
    return chunks

def extract_text_from_docx(path: str) -> str:
    try:
        doc = Document(path)
        return "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        logger.warning(f"DOCX text extraction failed for {path}: {e}")
        return ""

def split_sentences(text: str) -> List[str]:
    if NLP_MODEL:
        try: return [sent.text.strip() for sent in NLP_MODEL(text).sents if sent.text.strip()]
        except Exception: pass
    return [s.strip() for s in re.split(r'(?<=[.!?])\s+', text) if s.strip()]

# -------------------------
# ADGM DATA FETCHER
# -------------------------

class ADGMDataFetcher:
    def __init__(self, cache_dir: str = None):
        self.cache_dir = cache_dir or os.path.join(tempfile.gettempdir(), "adgm_cache")
        os.makedirs(self.cache_dir, exist_ok=True)
        self.session = requests.Session()
        self.session.headers.update({'User-Agent': 'Mozilla/5.0'})

    def fetch_document_text(self, url: str) -> Optional[str]:
        cache_path = os.path.join(self.cache_dir, f"{hash(url)}.txt")
        if os.path.exists(cache_path):
            try:
                with open(cache_path, 'r', encoding='utf-8') as f: return f.read()
            except Exception as e:
                logger.warning(f"Cache read failed for {url}: {e}")
        try:
            response = self.session.get(url, timeout=20, allow_redirects=True)
            response.raise_for_status()
            content = None
            if url.lower().endswith('.pdf') and PDF_PARSER_ENABLED:
                with fitz.open(stream=response.content, filetype="pdf") as doc: content = "".join(page.get_text() for page in doc)
            elif url.lower().endswith('.docx'):
                with open(os.path.join(self.cache_dir, "temp.docx"), 'wb') as f: f.write(response.content)
                content = extract_text_from_docx(os.path.join(self.cache_dir, "temp.docx"))
            else:
                content = re.sub('<[^<]+?>', '', response.text)
            if content:
                with open(cache_path, 'w', encoding='utf-8') as f: f.write(content)
            return content
        except Exception as e:
            logger.error(f"Failed to fetch or parse {url}: {e}")
            return None

    def build_knowledge_base(self) -> List[Dict]:
        documents = []
        for category_key, category_data in ADGM_DATA_SOURCES.items():
            for doc_key, url in category_data.get("documents", {}).items():
                content = self.fetch_document_text(url)
                if content: documents.append({"source": f"adgm_{category_key}_{doc_key}", "text": content, "citation": f"ADGM {category_data['name']} - {doc_key.replace('_', ' ').title()}"})
        logger.info(f"Built knowledge base with {len(documents)} real ADGM documents")
        return documents

# -------------------------
# CORE RAG SYSTEM
# -------------------------

class RAGSystem:
    def __init__(self, api_key: str, config: Config):
        self.api_key = api_key
        self.config = config
        self.client = None
        self.embedder = None
        self.index = None
        self.chunks = []
        self.metadata = []
        self.data_fetcher = ADGMDataFetcher()
        self.api_key_is_valid = True  # ** Assume key is valid until a call fails **

        if self.api_key:
            try:
                self.client = OpenAI(api_key=self.api_key)
            except Exception as e:
                logger.warning(f"OpenAI client initialization failed: {e}")
                self.client = None
                self.api_key_is_valid = False

        if 'SentenceTransformer' in globals() and SentenceTransformer:
            try:
                self.embedder = SentenceTransformer(self.config.embedding_model)
            except Exception as e:
                logger.warning(f"Embedder init failed: {e}")

    def _llm_call(self, prompt: str, max_tokens: int = None) -> str:
        if not self.client or not self.api_key_is_valid:
            return "LLM not available"
        try:
            response = self.client.chat.completions.create(
                model=self.config.llm_model, messages=[{"role": "user", "content": prompt}],
                temperature=self.config.temperature, max_tokens=max_tokens or self.config.max_tokens)
            return response.choices[0].message.content.strip()
        except Exception as e:
            # ** NEW: Detect key failure here and set the state flag **
            error_str = str(e).lower()
            if 'insufficient_quota' in error_str or 'authentication' in error_str or 'invalid_api_key' in error_str:
                logger.error(f"OpenAI API Key has failed: {e}. Disabling further AI calls.")
                self.api_key_is_valid = False
            return f"LLM call failed: {str(e)}"

    def build_index_from_real_data(self):
        if not self.embedder: return logger.warning("Embedder not available, skipping index build")
        documents = self.data_fetcher.build_knowledge_base()
        if not documents: return logger.warning("No ADGM documents fetched, RAG will be limited.")
        all_chunks, all_metadata = [], []
        for doc in documents:
            doc_chunks = chunk_text(doc.get('text', ''), self.config.chunk_size, self.config.chunk_overlap)
            for i, chunk in enumerate(doc_chunks):
                all_chunks.append(chunk['text'])
                all_metadata.append({'source': doc.get('source', ''), 'citation': doc.get('citation', ''), 'text': chunk['text'][:500], 'chunk_id': i })
        if all_chunks and 'faiss' in globals() and faiss:
            try:
                embeddings = self.embedder.encode(all_chunks, convert_to_numpy=True)
                self.index = faiss.IndexFlatL2(embeddings.shape[1]); self.index.add(embeddings.astype("float32"))
                self.chunks, self.metadata = all_chunks, all_metadata
                logger.info(f"Built RAG index with {len(all_chunks)} real ADGM chunks")
            except Exception as e: logger.error(f"Faiss index build failed: {e}")

    def retrieve(self, query: str, k: int = 3) -> List[Dict]:
        if not self.index: return []
        try:
            query_emb = self.embedder.encode([query], convert_to_numpy=True)
            _, indices = self.index.search(query_emb.astype("float32"), k)
            return [self.metadata[idx] for idx in indices[0] if idx < len(self.metadata)]
        except Exception as e:
            logger.error(f"Retrieval failed: {e}"); return []

    def analyze_compliance_chunk(self, text_chunk: str) -> Dict:
        # ** UPDATED: Check for valid client and valid key state **
        if not self.client or not self.api_key_is_valid:
            return {"issues": []}
        refs = self.retrieve(text_chunk)
        context = "\n\n".join([f"Reference from {ref.get('citation', '')}:\n{ref.get('text', '')}" for ref in refs])
        prompt = f"""Analyze the following text chunk for ADGM compliance based ONLY on the provided ADGM references.
**Text Chunk to Analyze:** "{text_chunk}"
**Official ADGM References:** {context}
Provide a JSON response with a list of issues. For each issue, include: "description", "regulation", "severity", and "fix".
If there are no issues, return an empty list in the JSON structure. Example: {{"issues": []}}
JSON Response:"""
        response = self._llm_call(prompt)
        try:
            return json.loads(response.strip().replace("```json", "").replace("```", ""))
        except json.JSONDecodeError:
            logger.warning(f"Failed to decode LLM JSON for chunk analysis: {response}"); return {"issues": []}
            
    def verify_ambiguous_language(self, sentence: str) -> bool:
        """Uses AI to verify if 'weak' language is genuinely ambiguous or acceptably permissive."""
        # ** UPDATED: Check for valid client and valid key state **
        if not self.client or not self.api_key_is_valid:
            return True  # Fallback: If AI fails, assume it's a real flag to be safe.

        prompt = f"""As a legal analysis expert, evaluate the following clause:
        **Clause:** "{sentence}"
        Is the use of a permissive term in this clause genuinely ambiguous and likely to cause legal uncertainty, or is it a standard grant of discretion?
        Respond with only 'YES' if it is ambiguous, or 'NO' if it is acceptable."""
        response = self._llm_call(prompt, max_tokens=10)
        return response.strip().upper() == 'YES'

# -------------------------
# ENHANCED DOCUMENT PROCESSOR
# -------------------------

class DocumentProcessor:
    def __init__(self, rag: RAGSystem, config: Config):
        self.rag = rag
        self.config = config
        self.checklists = ADGM_CHECKLISTS
        self.adgm_rules = {
            "jurisdiction": {"pattern": r"jurisdiction of the courts of (?!(abu dhabi global market|adgm))", "issue": "Incorrect jurisdiction specified. Must be ADGM Courts.", "severity": "High", "citation": "ADGM Courts, Civil Procedure Rules 2016", "source": "ADGM Rule"},
            "working_hours": {"pattern": r"work(?:ing)?\s+hours\s+exceed\s+(\d+)", "check": lambda m: int(m.group(1)) > 48, "issue": "Working hours exceed the 48-hour weekly limit.", "severity": "High", "citation": "ADGM Employment Regulations 2019, Art. 12", "source": "ADGM Rule"},
            "weak_language": {"pattern": r"\b(may|might|could|should consider)\b", "issue": "Ambiguous or non-binding language used. Use definitive terms like 'shall' or 'must'.", "severity": "Medium", "citation": "Contractual Best Practice", "source": "Pattern Match"}
        }

    def detect_document_type(self, text: str, filename: str = "") -> str:
        content = (text + " " + filename).lower()
        doc_type_indicators = {"articles of association": "Articles of Association", "employment contract": "Employment Contract", "compliance policy": "Compliance Policy Manual"}
        for indicator, doc_type in doc_type_indicators.items():
            if indicator in content: return doc_type
        return "Unknown Document"

    def detect_process_type(self, documents: List[Dict]) -> tuple:
        doc_types = [doc.get('document_type', '') for doc in documents]
        process_scores = {key: sum(1 for dt in doc_types if any(rd.lower() in dt.lower() for rd in data['required_documents'])) for key, data in self.checklists.items()}
        best_process = max(process_scores, key=process_scores.get) if any(v > 0 for v in process_scores.values()) else "general"
        process_info = self.checklists.get(best_process, {})
        required_docs = process_info.get('required_documents', [])
        found_docs = [rd for rd in required_docs if any(rd.lower() in dt.lower() for dt in doc_types)]
        missing_docs = [rd for rd in required_docs if rd not in found_docs]
        return best_process, process_info, found_docs, missing_docs

    def detect_red_flags_with_real_data(self, text: str) -> List[Dict]:
        flags = []
        # ** UPDATED: Check for valid key state before attempting AI analysis **
        if self.rag.client and self.rag.index and self.rag.api_key_is_valid:
            doc_chunks = chunk_text(text, self.config.chunk_size, self.config.chunk_overlap)
            for chunk in doc_chunks:
                # This will now safely return empty if the key fails mid-process
                analysis = self.rag.analyze_compliance_chunk(chunk['text'])
                for issue in analysis.get('issues', []):
                    flags.append({"category": issue.get('regulation', 'AI Compliance'), "issue": issue.get('description', ''), "severity": issue.get('severity', 'Medium'), "sentence": chunk['text'], "confidence": 0.85, "justification": f"Detected by AI based on official ADGM data. Reference: {issue.get('regulation', 'N/A')}", "fix": issue.get('fix', 'Review against ADGM regulations.'), "citation": issue.get('regulation', 'ADGM Regulations'), "source": "RAG Analysis"})
        
        for sentence in split_sentences(text):
            for rule_name, rule_data in self.adgm_rules.items():
                match = re.search(rule_data["pattern"], sentence, re.IGNORECASE)
                if match and rule_data.get("check", lambda m: True)(match):
                    is_real_flag = True
                    justification = f"Violates rule: {rule_data['citation']}"
                    source = rule_data["source"]
                    
                    if rule_name == 'weak_language':
                        # ** UPDATED: Check for valid key state before AI verification **
                        if self.rag.client and self.rag.api_key_is_valid:
                            logger.info(f"Verifying potential weak language: '{sentence[:80]}...'")
                            is_real_flag = self.rag.verify_ambiguous_language(sentence)
                            if is_real_flag:
                                justification = "AI confirmed this language is contextually ambiguous."
                                source = "AI-Verified Rule"
                        # If no client or key is invalid, it falls through and flags based on the pattern
                        # ensuring rules are not suppressed.

                    if is_real_flag:
                        flags.append({"category": rule_name.replace("_", " ").title(), "issue": rule_data["issue"], "severity": rule_data["severity"], "sentence": sentence, "confidence": 0.95, "justification": justification, "fix": "Update clause to comply with the cited ADGM regulation.", "citation": rule_data["citation"], "source": source})
        return flags

    def markup_docx_inline(self, path: str, flags: List[Dict]) -> str:
        try: doc = Document(path)
        except Exception as e: logger.error(f"Failed to open document: {e}"); return path
        if doc.paragraphs:
            header_para = doc.paragraphs[0].insert_paragraph_before()
            avg_conf = np.mean([f.get('confidence', 0.5) for f in flags]) if flags else 0.0
            rag_status = "✓ AI-Powered" if self.rag.client and self.rag.api_key_is_valid else "✗ Rules-Only"
            header_text = f"[ADGM COMPLIANCE ANALYSIS | Mode: {rag_status} | Flags: {len(flags)} | Avg. Confidence: {avg_conf:.2f}]"
            run = header_para.add_run(header_text)
            run.font.name = 'Calibri'; run.font.size = Pt(10); run.font.color.rgb = RGBColor(255, 0, 0); run.bold = True
        for i, flag in enumerate(flags, 1):
            sentence_to_find = flag.get("sentence", "")[:200].lower()
            for para in doc.paragraphs:
                if sentence_to_find and sentence_to_find in para.text.lower():
                    self._highlight_paragraph_enhanced(para, flag.get('severity', 'Medium'))
                    source = flag.get("source", "Unknown"); color = {"RAG Analysis": RGBColor(0, 100, 0), "ADGM Rule": RGBColor(255, 0, 0), "AI-Verified Rule": RGBColor(255, 140, 0)}.get(source, RGBColor(0, 0, 255)); prefix = {"RAG Analysis": "🔍 AI-DETECTED", "ADGM Rule": "⚖️ REGULATION", "AI-Verified Rule": "🤖 AI-VERIFIED"}.get(source, "🔎 PATTERN")
                    comment_run = para.add_run(f"\n[{prefix} #{i} | {flag.get('severity')}]: {flag.get('issue', '')} (Citation: {flag.get('citation')})")
                    comment_run.font.name = 'Arial'; comment_run.font.size = Pt(9); comment_run.font.color.rgb = color; comment_run.italic = True
                    break
        base_name = os.path.splitext(os.path.basename(path))[0]
        out_path = os.path.join(tempfile.gettempdir(), f"{base_name}_REVIEWED_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
        try:
            doc.save(out_path); logger.info(f"Saved enhanced marked-up document: {out_path}"); return out_path
        except Exception as e:
            logger.error(f"Save failed: {e}"); return path

    def _highlight_paragraph_enhanced(self, paragraph, severity: str):
        color_map = {"High": WD_COLOR_INDEX.RED, "Medium": WD_COLOR_INDEX.YELLOW, "Low": WD_COLOR_INDEX.TURQUOISE}
        highlight_color = color_map.get(severity, WD_COLOR_INDEX.GRAY_25)
        for run in paragraph.runs: run.font.highlight_color = highlight_color

# -------------------------
# ENHANCED ADGM AGENT
# -------------------------

class ADGMAgent:
    def __init__(self, api_key: str = None):
        self.config = Config()
        self.rag = RAGSystem(api_key, self.config)
        self.processor = DocumentProcessor(self.rag, self.config)
        self._initialize_with_real_data()

    def _initialize_with_real_data(self):
        logger.info("Initializing ADGM Agent with real data sources...")
        try:
            self.rag.build_index_from_real_data()
            logger.info("✓ Successfully initialized with real ADGM data sources")
        except Exception as e:
            logger.warning(f"⚠️ Failed to load real ADGM data: {e}. Falling back to limited mode.")

    def analyze_docx_files(self, file_paths: List[str]) -> Dict:
        if not file_paths: return {"error": "No files provided"}
        docx_files = [f for f in file_paths if f.lower().endswith('.docx')]
        if not docx_files: return {"error": "No .docx files found. Only DOCX files are supported."}
        total_size_bytes = sum(os.path.getsize(p) for p in docx_files if os.path.exists(p))
        self.config.adjust_for_upload_size(total_size_bytes)
        documents, results, marked_up_files = [], [], []
        for path in docx_files:
            text = extract_text_from_docx(path)
            if not text: continue
            documents.append({"text": text, "filename": os.path.basename(path), "document_type": self.processor.detect_document_type(text, os.path.basename(path)), "path": path })
        if not documents: return {"error": "No valid DOCX files could be processed"}
        process_type, process_info, found_docs, missing_docs = self.processor.detect_process_type(documents)
        total_flags = 0
        for doc_data in documents:
            flags = self.processor.detect_red_flags_with_real_data(doc_data["text"])
            total_flags += len(flags)
            marked_up_path = self.processor.markup_docx_inline(doc_data["path"], flags)
            marked_up_files.append(marked_up_path)
            results.append({"file": doc_data["filename"], "document_type": doc_data["document_type"], "status": "processed", "flags": flags, "marked_up_file": marked_up_path })
        bundle_path = self._create_enhanced_bundle(marked_up_files, results, process_info, found_docs, missing_docs, total_flags)
        return {"analysis_id": str(uuid.uuid4()), "timestamp": datetime.now().isoformat(), "process": process_info.get('name', process_type.replace('_', ' ').title()), "completeness_status": "Complete" if not missing_docs else f"Missing {len(missing_docs)} documents", "found_documents": found_docs, "missing_documents": missing_docs, "total_issues": total_flags, "documents": results, "bundle": bundle_path}

    def _create_enhanced_bundle(self, marked_up_files, results, process_info, found_docs, missing_docs, total_flags) -> str:
        if not marked_up_files: return None
        bundle_name = f"ADGM_Analysis_Bundle_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
        bundle_path = os.path.join(tempfile.gettempdir(), bundle_name)
        try:
            with zipfile.ZipFile(bundle_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for file_path in marked_up_files:
                    if os.path.exists(file_path): zipf.write(file_path, os.path.basename(file_path))
                report_content = self._generate_enhanced_report(results, process_info, found_docs, missing_docs, total_flags)
                zipf.writestr("ADGM_Compliance_Report.txt", report_content)
            logger.info(f"Created enhanced bundle: {bundle_path}"); return bundle_path
        except Exception as e:
            logger.error(f"Enhanced bundle creation failed: {e}"); return None

    def _generate_enhanced_report(self, results, process_info, found_docs, missing_docs, total_flags) -> str:
        data_quality = "HIGH (Real ADGM Data)" if self.rag.index else "MEDIUM (Pattern-based)"
        report = f"""ADGM COMPLIANCE ANALYSIS REPORT\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | Data Quality: {data_quality}\n{'='*80}\nEXECUTIVE SUMMARY:\nProcess Type: {process_info.get('name', 'Unknown')} | Total Issues Found: {total_flags}\n\nDOCUMENT COMPLETENESS:\n{'✓ FOUND:' if found_docs else ''}{chr(10).join(f'  • {doc}' for doc in found_docs)}\n{'✗ MISSING:' if missing_docs else ''}{chr(10).join(f'  • {doc}' for doc in missing_docs)}\n{'='*80}"""
        for i, result in enumerate(results, 1):
            report += f"\nDOCUMENT {i}: {result.get('file', 'Unknown')} ({result.get('document_type', 'Unknown')})\n"
            for j, flag in enumerate(result.get('flags', []), 1):
                report += f"""  - Issue #{j} ({flag.get('severity')}, Source: {flag.get('source')}): {flag.get('issue')}\n    Justification: {flag.get('justification')}\n"""
        return report

# -------------------------
# MAIN APPLICATION (Gradio UI)
# -------------------------

if __name__ == "__main__":
    def analyze_files_interface(api_key, files):
        if not files:
            return "Please upload at least one DOCX file for analysis.", None, None
        
        clean_api_key = api_key.strip() if api_key else None
        agent = ADGMAgent(api_key=clean_api_key)
        paths = [f.name for f in files]
        result = agent.analyze_docx_files(paths)
        
        if "error" in result:
            return result["error"], None, None
            
        json_report_path = os.path.join(tempfile.gettempdir(), f"{result['analysis_id']}.json")
        with open(json_report_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, indent=2, ensure_ascii=False)

        bundle_path = result.get("bundle")
        
        # ** UPDATED: UI reporting is now based on the reliable state flag **
        mode_message = ""
        if clean_api_key:
            if agent.rag.api_key_is_valid:
                mode_message = "**Mode:** Hybrid Mode (AI + Rules)"
            else:
                mode_message = "**Mode:** <span style='color:red;'>API Key failed. Falling back to Rules-Only Mode.</span>"
        else:
            mode_message = "**Mode:** Rules-Only Mode (No API Key provided)"

        summary = f"""✅ **Analysis Complete**
- {mode_message}
- **Process Detected:** {result.get('process')}
- **Total Issues Found:** {result.get('total_issues')}
- **Completeness:** {result.get('completeness_status')}
**Document Checklist:**
{chr(10).join(f"  ✓ {doc}" for doc in result.get('found_documents', []))}
{chr(10).join(f"  ✗ {doc}" for doc in result.get('missing_documents', []))}
**Next Steps:**
1. Download the 'Marked-up DOCX Bundle' below.
2. Open the DOCX files to see issues highlighted directly in your documents.
3. Refer to the JSON Report for a full, detailed breakdown."""
        return summary, json_report_path, bundle_path

    ui = gr.Interface(
        fn=analyze_files_interface,
        inputs=[
            gr.Textbox(type="password", label="OpenAI API Key (Optional)", placeholder="Enter your key for AI-enhanced analysis (sk-...)"),
            gr.File(file_types=[".docx"], file_count="multiple", label="Upload ADGM Corporate Documents (.docx)")
        ],
        outputs=[
            gr.Markdown(label="Analysis Summary"),
            gr.File(label="📊 Download Full JSON Report"),
            gr.File(label="📁 Download Marked-up DOCX Bundle (.zip)")
        ],
        title="🏢 ADGM Corporate Compliance Agent",
        description="""**Analyzes corporate documents against official ADGM regulations.** This tool uses a hybrid approach: high-precision rules and an advanced AI (RAG) system. Highly recommended: Providing a valid OpenAI API key enables AI-verification for nuanced issues, relying on regex mode for ambiguous language may lead to several false positives.""",
        article="""*Disclaimer: This tool provides guidance and is not a substitute for professional legal advice. Always consult with qualified legal counsel.*""",
        theme=gr.themes.Soft(),
        allow_flagging="never"
    )

    ui.launch(server_name="0.0.0.0", server_port=7860)